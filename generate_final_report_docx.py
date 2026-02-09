from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_professional_report():
    doc = Document()

    # Style definitions
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Title Page
    title = doc.add_heading('Final Project Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Change Point Analysis of Brent Oil Prices (1987-2022)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph('\n' * 2)
    author = doc.add_paragraph('Developed by: Antigravity AI\nDate: February 9, 2026')
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "This project investigates the historical price dynamics of Brent Crude Oil from May 1987 to September 2022. "
        "The primary goal was to identify significant structural breaks (change points) and quantify the association "
        "between these shifts and major geopolitical or economic events. Using Bayesian inference via PyMC, we "
        "successfully identified a major regime shift in 2005, transitioning the market into a sustained higher-price environment."
    )

    # 2. Methodology
    doc.add_heading('2. Methodology', level=1)
    doc.add_paragraph(
        "Our analysis followed a rigorous data science workflow:\n"
    )
    methodology = [
        "Data Collection: Ingestion of historical spot prices from the U.S. Energy Information Administration (EIA).",
        "Exploratory Data Analysis (EDA): Assessment of stationarity (ADF Test) and volatility clustering.",
        "Bayesian Modeling: Implementation of a Change Point model with discrete switch point (tau) priors.",
        "Interactive Visualization: Development of a full-stack dashboard (Flask/React) for stakeholder exploration."
    ]
    for item in methodology:
        doc.add_paragraph(item, style='List Bullet')

    # 3. Statistical Analysis (EDA)
    doc.add_heading('3. Statistical Analysis Results', level=1)
    doc.add_paragraph(
        "Initial testing revealed that Brent oil prices are non-stationary in their raw form but become "
        "stationary when converted to log returns. This justifies the use of models designed to capture "
        "shifts in average levels or trends."
    )
    
    # Insert EDA Plot
    if os.path.exists('results/brent_prices.png'):
        doc.add_paragraph('Figure 1: Historical Brent Oil Price Trend (1987-2022)', style='Caption')
        doc.add_picture('results/brent_prices.png', width=Inches(6))

    # 4. Bayesian Change Point Results
    doc.add_heading('4. Bayesian Change Point Modeling', level=1)
    doc.add_paragraph(
        "The Bayesian model successfully converged (R-hat = 1.02) and identified a definitive status "
        "change in the market structure."
    )

    # Table of Results
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'

    results_data = [
        ('Detected Change Point Date', 'March 02, 2005'),
        ('Regime 1 Mean Price', '$21.46'),
        ('Regime 2 Mean Price', '$75.90'),
        ('Relative Price Increase', '253.7%'),
        ('Model Certainty', 'High (Sharp Posterior Peak)')
    ]

    for metric, value in results_data:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = value

    doc.add_paragraph('\n')
    if os.path.exists('results/fitted_model.png'):
        doc.add_paragraph('Figure 2: Identified Change Point and Regime Means', style='Caption')
        doc.add_picture('results/fitted_model.png', width=Inches(6))

    # 5. Geopolitical Event Association
    doc.add_heading('5. Geopolitical Event Association', level=1)
    doc.add_paragraph(
        "The identified March 2005 change point marks the beginning of the mid-2000s energy super-cycle. "
        "While specific events like the Iraq War (2003) created immediate volatility, the 2005 shift represents "
        "a fundamental structural realignment caused by massive industrial growth in emerging markets (China/India) "
        "and tightening global production capacity."
    )

    # 6. Conclusion
    doc.add_heading('6. Conclusion', level=1)
    doc.add_paragraph(
        "The project provides a transparent, probabilistic framework for understanding market regimes. "
        "The 253% increase identified in 2005 highlights how long-term structural changes far outweigh "
        "the impact of isolated short-term events. The provided interactive dashboard remains available "
        "for further granular exploration of current and future market shocks."
    )

    filename = 'Final_Project_Report_CPA.docx'
    doc.save(filename)
    print(f"Professional report generated: {filename}")

if __name__ == "__main__":
    create_professional_report()

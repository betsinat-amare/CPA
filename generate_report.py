from docx import Document
from docx.shared import Inches
import pandas as pd
import os

def create_report():
    doc = Document()
    doc.add_heading('Interim Report: Change Point Analysis of Brent Oil Prices', 0)

    # Section 1: Planned Analysis Steps
    doc.add_heading('1. Planned Analysis Steps', level=1)
    steps = [
        "Data Ingestion & Preprocessing: Loading historical Brent oil price data (1987-2022) and cleaning for missing values.",
        "Exploratory Data Analysis (EDA): Analyzing trends, seasonality, stationarity (ADF test), and volatility clustering.",
        "Event Research: Identifying 10-15 major geopolitical/economic events that impacted the oil market.",
        "Bayesian Change Point Modeling: Using PyMC to detect structural breaks in the price series.",
        "Impact Quantification: Measuring the magnitude of price shifts associated with specific events.",
        "Communication & Dashboard: Visualizing findings in an interactive Streamlit dashboard."
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Bullet')

    # Section 2: Key Geopolitical Events
    doc.add_heading('2. Structured Event Dataset', level=1)
    doc.add_paragraph("The following table summarizes the key events identified for the analysis:")
    
    events_data = [
        ["Date", "Event", "Type"],
        ["20-May-87", "Start of Dataset (Baseline)", "Data Baseline"],
        ["02-Aug-90", "Invasion of Kuwait by Iraq", "Geopolitical"],
        ["17-Jan-91", "Start of Gulf War", "Geopolitical"],
        ["01-Jul-97", "Asian Financial Crisis", "Economic"],
        ["11-Sep-01", "September 11 Attacks", "Geopolitical"],
        ["20-Mar-03", "US-led Invasion of Iraq", "Geopolitical"],
        ["15-Sep-08", "Lehman Brothers Collapse", "Economic"],
        ["15-Feb-11", "Arab Spring Beginnings", "Geopolitical"],
        ["01-Jun-14", "Oil Price Crash (Shale Revolution)", "Economic"],
        ["30-Nov-16", "OPEC+ Production Cut Agreement", "Policy"],
        ["11-Mar-20", "WHO declares COVID-19 pandemic", "Health/Economic"],
        ["24-Feb-22", "Russian Invasion of Ukraine", "Geopolitical"]
    ]
    
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Event'
    hdr_cells[2].text = 'Type'
    
    for date, event, etype in events_data[1:]:
        row_cells = table.add_row().cells
        row_cells[0].text = date
        row_cells[1].text = event
        row_cells[2].text = etype

    # Section 3: Initial EDA Findings
    doc.add_heading('3. Initial EDA Findings', level=1)
    doc.add_paragraph("The preliminary analysis of Brent oil prices (1987-2022) reveals several key characteristics:")
    
    eda_points = [
        "Non-Stationarity: Raw price data shows a strong stochastic trend, failing the ADF test at the level.",
        "Volatility Clustering: Significant periods of high volatility are observed during crises (e.g., 2008, 2020).",
        "Regime Shifts: Distinct levels of mean prices are visible, shifting from $20 range (pre-2000s) to over eth $100 range during peak periods.",
        "Returns Distribution: Price returns exhibit heavy tails, suggesting that 'black swan' events have a significant impact on the market."
    ]
    for point in eda_points:
        doc.add_paragraph(point, style='List Bullet')

    doc.save('Interim_Report_Task_1.docx')
    print("Report generated successfully: Interim_Report_Task_1.docx")

if __name__ == "__main__":
    create_report()
